"""
Document processor for extracting content from various file formats
Designed for minimal performance impact and graceful error handling
"""

import os
from typing import List, Tuple, Dict, Any
from pathlib import Path
import logging
from io import BytesIO

from data.document_models import Document, DocumentType, DocumentStatus

# Import document processing libraries with graceful fallbacks
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_bytes
    # Configure Tesseract path for Windows
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class DocumentProcessor:
    """Handles document content extraction with minimal system impact"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
    def process_document(self, file_content: bytes, filename: str, content_type: str) -> Tuple[Document, List[str]]:
        """Process a document and extract its content"""
        
        # Create document object with required fields
        document = Document(
            title=filename,
            filename=filename,
            original_filename=filename,
            file_path="",  # Will be set by storage manager
            mime_type=content_type,
            file_size=len(file_content),
            status=DocumentStatus.PROCESSING
        )
        
        try:
            # Extract content based on file type
            if content_type == "application/pdf" and PDF_AVAILABLE:
                content, chunks = self._extract_pdf(file_content, filename)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" and DOCX_AVAILABLE:
                content, chunks = self._extract_docx(file_content, filename)
            elif content_type == "text/plain":
                content, chunks = self._extract_text(file_content, filename)
            else:
                # Fallback for unsupported types
                content, chunks = self._extract_fallback(file_content, filename)
            
            # Set document content and metadata
            document.content = content
            document.document_type = self._classify_document(content, filename)
            document.summary = self._generate_summary(content)
            document.key_insights = self._extract_key_insights(content, document.document_type)
            document.status = DocumentStatus.INDEXED
            
            return document, chunks
            
        except Exception as e:
            self.logger.error(f"Error processing document {filename}: {e}")
            document.status = DocumentStatus.ERROR
            document.metadata = {"error": str(e), "error_type": type(e).__name__}
            # Print error for debugging
            print(f"Document processing error for {filename}: {e}")
            return document, []

    def _extract_pdf(self, file_content: bytes, filename: str) -> Tuple[str, List[str]]:
        """Extract text from PDF using PyPDF2, with OCR fallback for image-based PDFs"""
        try:
            pdf_stream = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            text_content = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
            
            pdf_stream.close()
            
            full_text = "\n\n".join(text_content)
            
            # If no text was extracted or very little text, try OCR
            if len(full_text.strip()) < 50 and OCR_AVAILABLE:
                self.logger.info(f"PDF {filename} appears to be image-based, attempting OCR...")
                print(f"PDF {filename} appears to be image-based, attempting OCR...")
                ocr_text = self._extract_pdf_with_ocr(file_content, filename)
                if len(ocr_text.strip()) > len(full_text.strip()):
                    full_text = ocr_text
                    print(f"OCR extracted {len(ocr_text)} characters from {filename}")
            
            chunks = self._create_chunks(full_text)
            
            return full_text, chunks
            
        except Exception as e:
            self.logger.error(f"Error extracting PDF {filename}: {e}")
            raise

    def _extract_pdf_with_ocr(self, file_content: bytes, filename: str) -> str:
        """Extract text from image-based PDF using OCR"""
        try:
            if not OCR_AVAILABLE:
                return ""
            
            # Convert PDF pages to images with poppler path
            poppler_path = r"C:\poppler\poppler-23.08.0\Library\bin"
            images = convert_from_bytes(
                file_content, 
                dpi=300,
                poppler_path=poppler_path
            )
            
            ocr_text = []
            for i, image in enumerate(images):
                # Use pytesseract to extract text from image
                page_text = pytesseract.image_to_string(image, lang='eng')
                if page_text.strip():
                    # Clean up OCR artifacts
                    cleaned_text = self._clean_ocr_text(page_text)
                    if cleaned_text.strip():
                        ocr_text.append(f"--- Page {i+1} ---\n{cleaned_text}")
            
            return "\n\n".join(ocr_text)
            
        except Exception as e:
            self.logger.error(f"OCR extraction failed for {filename}: {e}")
            print(f"OCR extraction failed for {filename}: {e}")
            return ""

    def _extract_docx(self, file_content: bytes, filename: str) -> Tuple[str, List[str]]:
        """Extract text from DOCX using BytesIO"""
        try:
            docx_stream = BytesIO(file_content)
            doc = DocxDocument(docx_stream)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            docx_stream.close()
            
            full_text = "\n".join(text_content)
            chunks = self._create_chunks(full_text)
            
            return full_text, chunks
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX {filename}: {e}")
            raise

    def _extract_text(self, file_content: bytes, filename: str) -> Tuple[str, List[str]]:
        """Extract text from plain text file"""
        try:
            # Try UTF-8 first, fallback gracefully
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text = file_content.decode('latin-1')
                except UnicodeDecodeError:
                    text = file_content.decode('utf-8', errors='ignore')
            
            chunks = self._create_chunks(text)
            return text, chunks
            
        except Exception as e:
            self.logger.error(f"Error extracting text {filename}: {e}")
            raise

    def _extract_fallback(self, file_content: bytes, filename: str) -> Tuple[str, List[str]]:
        """Fallback extraction for unsupported file types"""
        try:
            # Try to extract as text
            text = file_content.decode('utf-8', errors='ignore')
            if not text.strip():
                text = f"Binary file: {filename} (Size: {len(file_content)} bytes)"
            
            chunks = self._create_chunks(text)
            return text, chunks
        except:
            text = f"File: {filename} (Size: {len(file_content)} bytes) - Content extraction not supported"
            return text, [text]

    def _create_chunks(self, text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
        """Split text into overlapping chunks for vector storage"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence or paragraph boundaries
            if end < len(text):
                for i in range(end - 50, end):
                    if i > start and text[i] in '.!?\n':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = max(start + chunk_size - overlap, start + 1)
            if start >= len(text):
                break
        
        return chunks

    def _classify_document(self, content: str, filename: str) -> DocumentType:
        """Classify document based on content and filename"""
        content_lower = content.lower()
        filename_lower = filename.lower()
        
        # Simple keyword-based classification
        if any(word in content_lower or word in filename_lower 
               for word in ['certificate', 'award', 'achievement', 'diploma']):
            return DocumentType.ACHIEVEMENT
        
        if any(word in content_lower or word in filename_lower 
               for word in ['project', 'implementation', 'development']):
            return DocumentType.PROJECT
        
        if any(word in content_lower or word in filename_lower 
               for word in ['skill', 'competency', 'training']):
            return DocumentType.SKILL
        
        if any(word in content_lower or word in filename_lower 
               for word in ['meeting', 'notes', 'minutes']):
            return DocumentType.MEETING_NOTES
        
        if any(word in content_lower or word in filename_lower 
               for word in ['reference', 'manual', 'guide']):
            return DocumentType.REFERENCE
        
        return DocumentType.GENERAL

    def _generate_summary(self, content: str, max_length: int = 200) -> str:
        """Generate a simple summary of the document content"""
        if len(content) <= max_length:
            return content
        
        # Extract first few sentences
        sentences = content.replace('\n', ' ').split('.')
        summary = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(summary + sentence + ".") <= max_length:
                summary += sentence + "."
            else:
                break
        
        if not summary:
            summary = content[:max_length] + "..."
        
        return summary.strip()

    def _clean_ocr_text(self, text: str) -> str:
        """Clean up OCR artifacts and encoding issues"""
        import re
        
        # Remove common OCR artifacts
        text = text.replace('ÿþ', '')  # BOM characters
        text = text.replace('ep CAMBRIDGE —', 'CAMBRIDGE')  # OCR misreading
        text = text.replace('—', '-')  # Replace em-dash with regular dash
        text = text.replace('\u2019', "'")  # Replace curly apostrophe
        text = text.replace('\u201c', '"').replace('\u201d', '"')  # Replace curly quotes
        
        # Clean up multiple spaces and normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove lines that are just punctuation or single characters
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if len(line) > 2 and not re.match(r'^[^\w]*$', line):
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines).strip()

    def _extract_key_insights(self, content: str, doc_type: DocumentType) -> List[str]:
        """Extract key insights based on document type"""
        insights = []
        content_lower = content.lower()
        
        # Simple keyword extraction based on document type
        if doc_type == DocumentType.PROJECT:
            tech_keywords = ['python', 'javascript', 'react', 'database', 'api', 'ml', 'ai']
            found_tech = [tech for tech in tech_keywords if tech in content_lower]
            if found_tech:
                insights.append(f"Technologies: {', '.join(found_tech)}")
        
        elif doc_type == DocumentType.ACHIEVEMENT:
            if 'certificate' in content_lower:
                insights.append("Contains certification information")
        
        return insights


# Global instance for reuse
_processor_instance = None

def get_document_processor() -> DocumentProcessor:
    """Get document processor instance (singleton pattern)"""
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = DocumentProcessor()
    return _processor_instance
